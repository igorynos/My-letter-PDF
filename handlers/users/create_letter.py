from aiogram.dispatcher.storage import FSMContext
from aiogram import types
from aiogram.types import CallbackQuery


from loader import dp, db, bot
from data.config import DATABASE

from states import State_list


from keyboards.inline.choise_to_letter import lst_template_to_letter, lst_cont_users_to_letter
from keyboards.inline.callback_data import lst_template_to_letter_callback, lst_cont_user_to_letter_callback
from keyboards.default.main_menu import main_menu


import docx
import io
import urllib.request
import re
from docx2pdf import convert
import aiohttp
from datetime import datetime
import platform
import subprocess

os_name = platform.system()


month_dict = {
    '01': 'января',
    '02': 'февраля',
    '03': 'марта',
    '04': 'апреля',
    '05': 'мая',
    '06': 'июня',
    '07': 'июля',
    '08': 'августа',
    '09': 'сентября',
    '10': 'октября',
    '11': 'ноября',
    '12': 'декабря'
}
current_datetime = datetime.now()

current_date = current_datetime.strftime('%Y-%m-%d')
doc_data = current_datetime.strftime(
    '%d %m %Y')
doc_data = doc_data.split()
doc_data[1] = month_dict[f'{doc_data[1]}']
doc_data = ' '.join(map(str, doc_data))


dict_temp_template = {}
dict_temp_cont_user = {}

dict_user_none_text = {}
dict_user_indx = {}


@dp.message_handler(text='Создать письмо')
async def choise_template_to_letter_1(message: types.Message):
    await lst_template_to_letter(message)


@dp.callback_query_handler(lst_template_to_letter_callback.filter())
async def choise_template_to_letter_2(call: CallbackQuery, callback_data: dict, state: FSMContext):
    result = callback_data.get('id')
    dict_temp_template[f'{call.message.chat.id}'] = result
    await lst_cont_users_to_letter(call.message)


@dp.callback_query_handler(lst_cont_user_to_letter_callback.filter())
async def choise_cont_users_to_letter_2(call: CallbackQuery, callback_data: dict, state: FSMContext):
    result = callback_data.get('cont_id')
    dict_temp_cont_user[f'{call.message.chat.id}'] = result
    await create_letter_1(call.message, state=state)


async def create_letter_1(message: types.Message, state: FSMContext):
    await message.answer(text=f'Пожалуйста подождите...')
    sql = f"SELECT link FROM template WHERE id = '{dict_temp_template[f'{message.chat.id}']}'"
    link_template = db.execute(sql, fetchone=True)
    link_template = link_template[0]

    response = urllib.request.urlopen(link_template)
    async with aiohttp.ClientSession() as session:
        async with session.get(link_template) as response:
            if response.status == 200:
                content = await response.read()

                with io.BytesIO(content) as f:
                    doc = docx.Document(f)

                    full_text = []
                    for paragraph in doc.paragraphs:
                        if '{{' in paragraph.text:
                            text = paragraph.text
                            pattern = r'\{\{(.+?)\}\}'
                            matches = re.findall(pattern, text)
                            full_text.extend(matches)

    sql_value = f"SELECT * FROM cont_users WHERE cont_id='{dict_temp_cont_user[f'{message.chat.id}']}'"
    sql_key = f"SELECT ORDINAL_POSITION, COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_SCHEMA = '{DATABASE}' AND TABLE_NAME = 'cont_users'"

    cont_card_value = db.execute(sql_value, fetchall=True, commit=True)
    cont_card_value = cont_card_value[0]
    cont_card_key = db.execute(sql_key, fetchall=True, commit=True)

    sql_value_my = f"SELECT * FROM users WHERE id='{message.chat.id}'"
    sql_key_my = f"SELECT ORDINAL_POSITION, COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_SCHEMA = '{DATABASE}' AND TABLE_NAME = 'users'"

    my_card_value = db.execute(sql_value_my, fetchall=True, commit=True)
    my_card_value = my_card_value[0]
    my_card_key = db.execute(sql_key_my, fetchall=True, commit=True)

    try:
        data = await state.get_data()
        dict_oper = data['dict_oper']
    except:
        dict_oper = {}
        for i, x in enumerate(cont_card_key):
            dict_oper[f"{x[1]}"] = cont_card_value[i]
        for i, x in enumerate(my_card_key):
            dict_oper[f"{x[1]}"] = my_card_value[i]
        if "data" in full_text:
            dict_oper[f"data"] = doc_data
        if "doc_text" in full_text:
            dict_oper[f"doc_text"] = ''

    lst_none = []

    for x in full_text:
        try:
            if dict_oper[x] == "":
                lst_none.append(x)
        except:
            lst_none.append(x)
    await state.update_data(
        {
            'dict_oper': dict_oper,
            'lst_none': lst_none,
            'full_text': full_text,
            'link_template': link_template
        }
    )
    if len(lst_none) != 0:
        await user_none(message=message, state=state)
    else:
        await create_letter_2(message=message, state=state)


async def create_letter_2(message: types.Message, state: FSMContext):
    await message.answer(text=f'Письмо создаётся...')
    data = await state.get_data()
    dict_oper = data['dict_oper']
    full_text = data['full_text']
    link_template = data['link_template']

    response = urllib.request.urlopen(link_template)
    async with aiohttp.ClientSession() as session:
        async with session.get(link_template) as response:
            if response.status == 200:
                content = await response.read()

                with io.BytesIO(content) as f:
                    doc = docx.Document(f)

                    sql = f"SELECT fio FROM users WHERE id='{message.chat.id}'"
                    name = db.execute(
                        sql, fetchall=True, commit=True)[0][0]

                    sql_1 = f"SELECT cont_org FROM cont_users WHERE cont_id='{dict_temp_cont_user[f'{message.chat.id}']}'"
                    cont_name = db.execute(
                        sql_1, fetchall=True, commit=True)[0][0]

                    for oper in full_text:
                        for paragraph in doc.paragraphs:
                            text = paragraph.text
                            search_str = f'{{{{{oper}}}}}'
                            paragraph.text = text.replace(
                                search_str, dict_oper[oper])
                            try:
                                paragraph.text = text.replace(
                                    {{{{{'data'}}}}}, doc_data)
                            except:
                                pass

                    dir_name = f"docs"
                    file_name = dir_name + \
                        f"/{name}_{cont_name}_{current_date}.docx".replace(
                            ' ', '_')
                    doc.save(file_name)
                    if os_name == 'Windows':
                        convert(file_name, file_name[:-5] + ".pdf")
                    elif os_name == 'Linux':
                        command = ['libreoffice', '--convert-to',
                                   'pdf', '--outdir', dir_name, file_name]
                        subprocess.run(command, check=True)
                    await bot.send_document(message.chat.id, open(file_name[:-5] + ".pdf", 'rb'), reply_markup=main_menu(message))


async def user_none(message: types.Message, state: FSMContext):
    data = await state.get_data()
    dict_user_none_text[f'{message.chat.id}'] = data['lst_none']
    dict_user_indx[f"{message.chat.id}"] = 0
    numb_0 = data['lst_none'][0]

    sql = f"SELECT full FROM lst_oper WHERE oper='{numb_0}'"
    result = db.execute(sql, fetchone=True, commit=True)
    result = result[0]

    await message.answer(text=f'Введите {result.lower()}')
    await State_list.first()


async def handle_answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    lst_none = data['lst_none']
    dict_oper = data['dict_oper']

    try:
        dict_oper[lst_none[dict_user_indx[f"{message.chat.id}"]]
                  ] = message.text
        try:
            dict_user_indx[f"{message.chat.id}"] += 1
            await state.update_data({'dict_oper': dict_oper, 'lst_none': lst_none})
            oper = lst_none[dict_user_indx[f"{message.chat.id}"]]
            result = db.execute(
                f"SELECT full FROM lst_oper WHERE oper='{oper}'", fetchone=True, commit=True)
            await message.answer(text=f'Введите {result[0].lower()}')
            await State_list.next()
        except:
            db.update_cont_user(
                cont_id=dict_temp_cont_user[f"{message.chat.id}"], dict_oper=dict_oper)
            db.update_user(id=message.chat.id, dict_oper=dict_oper)
            await create_letter_1(message=message, state=state)
    except:
        await state.finish()


@dp.message_handler(state=State_list.Q0)
@dp.message_handler(state=State_list.Q1)
@dp.message_handler(state=State_list.Q2)
@dp.message_handler(state=State_list.Q3)
@dp.message_handler(state=State_list.Q4)
@dp.message_handler(state=State_list.Q5)
@dp.message_handler(state=State_list.Q6)
@dp.message_handler(state=State_list.Q7)
@dp.message_handler(state=State_list.Q8)
@dp.message_handler(state=State_list.Q9)
@dp.message_handler(state=State_list.Q10)
@dp.message_handler(state=State_list.Q11)
@dp.message_handler(state=State_list.Q12)
@dp.message_handler(state=State_list.Q13)
@dp.message_handler(state=State_list.Q14)
@dp.message_handler(state=State_list.Q15)
@dp.message_handler(state=State_list.Q16)
@dp.message_handler(state=State_list.Q17)
@dp.message_handler(state=State_list.Q18)
@dp.message_handler(state=State_list.Q19)
async def answer_lst(message: types.Message, state: FSMContext):
    await handle_answer_lst(message, state)
